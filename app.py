import os
import io
import re
import PyPDF2
from docx import Document
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from transformers import pipeline
from youtube_transcript_api import YouTubeTranscriptApi
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

print("Loading AI model...")
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
print("AI Model Loaded!")

@app.route("/")
def home():
    return app.send_static_file("index.html")

@app.route("/<path:path>")
def static_files(path):
    return app.send_static_file(path)

# --------------------------
# TEXT EXTRACTION
# --------------------------

def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_youtube_text(url):
    patterns = [
        r"(?:v=)([A-Za-z0-9_-]{11})",
        r"(?:youtu\.be/)([A-Za-z0-9_-]{11})",
        r"(?:embed/)([A-Za-z0-9_-]{11})"
    ]

    video_id = None
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            video_id = match.group(1)
            break

    if not video_id:
        raise Exception("Invalid YouTube URL")

    transcript = YouTubeTranscriptApi.get_transcript(video_id)
    return " ".join([t["text"] for t in transcript])

# --------------------------
# AI SUMMARY
# --------------------------

def generate_summary(text):
    if len(text) > 1500:
        text = text[:1500]

    result = summarizer(text, max_length=150, min_length=40, do_sample=False)
    summary = result[0]["summary_text"]

    words = summary.replace(",", "").replace(".", "").split()
    keywords = list(set([w for w in words if len(w) > 4]))[:10]

    return summary, keywords

# --------------------------
# FILE CREATION
# --------------------------

def create_pdf(summary, keywords):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("<b>EduNote AI Summary</b>", styles["Heading1"]))
    elements.append(Paragraph(summary, styles["Normal"]))
    elements.append(Paragraph("<b>Key Topics:</b>", styles["Heading2"]))
    elements.append(Paragraph(", ".join(keywords), styles["Normal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def create_docx(summary, keywords):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("EduNote AI Summary", 0)
    doc.add_paragraph(summary)
    doc.add_heading("Key Topics", level=1)
    doc.add_paragraph(", ".join(keywords))
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --------------------------
# API ROUTES
# --------------------------

@app.route("/api/generate", methods=["POST"])
def generate_notes():
    try:
        input_type = request.form.get("type")

        if input_type == "text":
            text = request.form.get("text")

        elif input_type == "pdf":
            file = request.files.get("file")
            text = extract_text_from_pdf(file)

        elif input_type == "docx":
            file = request.files.get("file")
            text = extract_text_from_docx(file)

        elif input_type == "youtube":
            url = request.form.get("url")
            text = extract_youtube_text(url)

        else:
            return jsonify({"error": "Invalid input type"}), 400

        summary, keywords = generate_summary(text)

        return jsonify({
            "summary": summary,
            "keywords": keywords
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    data = request.json
    pdf_buffer = create_pdf(data["summary"], data["keywords"])
    return send_file(pdf_buffer,
                     as_attachment=True,
                     download_name="EduNote-Summary.pdf",
                     mimetype="application/pdf")

@app.route("/download/docx", methods=["POST"])
def download_docx():
    data = request.json
    docx_buffer = create_docx(data["summary"], data["keywords"])
    return send_file(docx_buffer,
                     as_attachment=True,
                     download_name="EduNote-Summary.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    print("Running on http://127.0.0.1:5000")
    app.run(debug=True)